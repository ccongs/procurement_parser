"""DOCX 문서 파서"""

import logging
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.table import Table

logger = logging.getLogger("docx_parser")


class DOCXParser:
    """DOCX 문서 파서"""

    def extract_text(self, file_path: Path) -> str:
        """전체 텍스트 추출"""
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"텍스트 추출 실패: {e}")
            return ""

    def extract_tables(self, file_path: Path) -> List[Dict[str, Any]]:
        """테이블 데이터 추출"""
        tables = []

        try:
            doc = Document(file_path)

            for i, table in enumerate(doc.tables):
                table_data = self._table_to_dict(table, i)
                if table_data:
                    tables.append(table_data)
        except Exception as e:
            logger.error(f"테이블 추출 실패: {e}")

        return tables

    def _table_to_dict(self, table: Table, index: int) -> Dict[str, Any]:
        """Table 객체를 딕셔너리로 변환"""
        rows = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            rows.append(row_data)

        if not rows:
            return {}

        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []

        return {
            "table_index": index,
            "headers": headers,
            "rows": data_rows,
            "raw_data": rows,
        }
